import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not found. Please install it using: pip install python-docx")
    sys.exit(1)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_docx():
    document = Document()

    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Campus ANPR System - Operation Manual"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.runs[0].font.name = 'Arial'
        header_para.runs[0].font.size = Pt(10)
        header_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Page ")
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(0, 0, 0)
        add_page_number(footer_run)
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.5 
    style.paragraph_format.first_line_indent = Inches(0.5) 
    
    def add_heading(text, level=1):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16 if level == 1 else (14 if level == 2 else 12))
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(24 if level == 1 else 16)
        p.paragraph_format.space_after = Pt(12)
            
    def add_screenshot_placeholder(text="[ INSERT SCREENSHOT OF INTERFACE HERE ]"):
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"\n{text}\n")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = True
        run.font.size = Pt(11)
        
    def add_bullet(text):
        p = document.add_paragraph(text)
        p.style = 'List Bullet'
        p.paragraph_format.first_line_indent = Inches(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(6)

    def add_code_block(text):
        # Splitting by newlines so each command is its own paragraph, avoiding weird justify stretching
        lines = text.split('\n')
        for line in lines:
            p = document.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    def add_normal_paragraph(text, justify=True):
        p = document.add_paragraph(text)
        if not justify:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
        return p

    # Title Page
    title_para = document.add_paragraph()
    title_para.paragraph_format.first_line_indent = Inches(0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    
    title_run = title_para.add_run("CAMPUS AUTOMATIC NUMBER PLATE RECOGNITION (ANPR) SYSTEM\n")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    subtitle_para = document.add_paragraph()
    subtitle_para.paragraph_format.first_line_indent = Inches(0)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("\nSYSTEM OPERATION & DEPLOYMENT MANUAL\n")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
    
    date_para = document.add_paragraph()
    date_para.paragraph_format.first_line_indent = Inches(0)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(200)
    date_run = date_para.add_run("Caraga State University - Cabadbaran City Campus")
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0, 0, 0)

    document.add_page_break()
    
    # Table of Contents
    add_heading('Table of Contents', level=1)
    
    page_header = document.add_paragraph()
    page_header.paragraph_format.first_line_indent = Inches(0)
    page_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ph = page_header.add_run("Page")
    run_ph.font.bold = True
    run_ph.font.name = 'Arial'
    run_ph.font.size = Pt(12)
    
    toc_items = [
        ("Title Page", "i", 0),
        ("Table of Contents", "ii", 0),
        ("Chapter I. System Overview", "1", 0),
        ("Chapter II. System Setup & Installation", "2", 0),
        ("2.1 System Compatibility & Hardware Requirements", "2", 1),
        ("2.2 Third-Party Software Installation", "3", 1),
        ("2.3 Installing Project Dependencies", "4", 1),
        ("2.4 Database Initialization & Schema Migration", "5", 1),
        ("2.5 Manual Setup Procedures", "6", 1),
        ("Chapter III. System Execution Procedures", "7", 0),
        ("Chapter IV. User Roles and Access Credentials", "8", 0),
        ("Chapter V. Core Module Functionality & Navigation", "9", 0),
        ("5.1 Administrator Module", "9", 1),
        ("5.2 Security Personnel Module", "10", 1),
        ("5.3 Standard User Self-Service Module", "11", 1),
        ("5.4 System Navigation Workflows", "12", 1),
        ("Chapter VI. Hardware & Camera Integration", "13", 0),
        ("Chapter VII. Troubleshooting & System Administration", "14", 0)
    ]
    
    for item, page, indent_level in toc_items:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(indent_level * 0.5)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        run = p.add_run(f"{item}\t{page}")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    document.add_page_break()

    # Chapter I
    add_heading('Chapter I. System Overview', level=1)
    add_normal_paragraph("The Campus Automatic Number Plate Recognition (ANPR) System is a comprehensive, full-stack application designed to secure and manage vehicle access at the Caraga State University - Cabadbaran City Campus. It utilizes advanced Optical Character Recognition (OCR) methodologies via Tesseract and PaddleOCR to read vehicle license plates and cross-reference them against a centralized database.")
    add_normal_paragraph("The system architecture consists of a FastAPI Python backend (Port 8000), a dedicated SMART-PLATE ANPR inference engine (Port 8003), and a highly responsive React/Vite frontend (Port 5173) which is deployed as an Electron Desktop Application for security guard checkpoints.")

    # Chapter II
    add_heading('Chapter II. System Setup & Installation', level=1)
    add_normal_paragraph("The project is configured for comprehensive deployment. The system relies on underlying hardware binaries and software frameworks that must be configured first.")
    
    add_heading('2.1 System Compatibility & Hardware Requirements', level=2)
    add_normal_paragraph("Portability: The system is completely portable. The entire repository can be installed and run on any directory or drive of any compliant laptop or desktop.")
    add_normal_paragraph("Operating System: The included automated batch scripts (.bat) are natively designed for Windows 10/11. However, the system's core codebase (Python/Node.js) is fully cross-platform and compatible with macOS and Linux distributions via manual terminal commands.")
    add_normal_paragraph("Hardware Specifications:", justify=False)
    add_bullet("Minimum RAM: 8GB System Memory.")
    add_bullet("Recommended RAM: 16GB System Memory (Required for stable SMART-PLATE OCR inference).")
    add_bullet("Processor: Multi-core CPU (Intel i5/Ryzen 5 or higher recommended).")
    
    add_heading('2.2 Third-Party Software Installation', level=2)
    add_normal_paragraph("Before initializing the software, several foundational applications must be installed globally on the host machine:")
    add_normal_paragraph("PostgreSQL Installation: Install PostgreSQL 14 or higher. During installation, set a master password for the 'postgres' superuser. Ensure that the PostgreSQL bin directory is added to the system PATH variables so that the backend can establish robust TCP/IP connections via SQLAlchemy.")
    add_normal_paragraph("Tesseract OCR Installation: The backend utilizes pytesseract for legacy plate recognition. Download the official Windows installer for Tesseract OCR. Install the binary precisely in C:\\Program Files\\Tesseract-OCR. This directory must be appended to the Windows System Environment Variables to prevent FileNotFoundError exceptions during execution.")
    add_normal_paragraph("PaddleOCR Configuration: PaddleOCR provides advanced neural-network-based character segmentation. It is installed automatically via Python PIP (paddlepaddle, paddleocr), however, it requires the Microsoft Visual C++ Redistributable 2015-2022 to be installed on the host machine to compile its C-based tensor operations.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: Tesseract Path Configuration in Environment Variables ]")

    add_heading('2.3 Installing Project Dependencies', level=2)
    add_normal_paragraph("Execute the 1_Install_Dependencies.bat located in the main directory. The script will automatically establish a Python virtual environment, install backend requirements via pip, and install frontend node modules via npm. Await the \"INSTALLATION COMPLETE!\" prompt before proceeding.")
    
    add_heading('2.4 Database Initialization & Schema Migration', level=2)
    add_normal_paragraph("Execute the 2_Setup_Database.bat file. When prompted, securely input your local PostgreSQL password.")
    add_normal_paragraph("Schema Migration: Under the hood, the system executes the seed_db.py Python module. This script utilizes SQLAlchemy's declarative base (Base.metadata.create_all(bind=engine)) to dynamically execute Data Definition Language (DDL) statements. It autonomously generates the relational SQL schemas, creating strongly-typed tables for Users, Vehicles, EntryLogs, Alerts, and Cameras.")
    add_normal_paragraph("Data Migration: Once the schema is constructed, the script performs an automated data migration, inserting hundreds of mock telemetry records, hashed administrative accounts, and baseline configuration flags to ensure the system is immediately usable for testing.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: pgAdmin Showing the Generated SQL Schema and Tables ]")

    add_heading('2.5 Manual Setup Procedures', level=2)
    add_normal_paragraph("For custom environments, the manual bootstrapping sequence is as follows:")
    
    add_normal_paragraph("1. Backend Environment Setup: Navigate to the backend directory, instantiate a virtual environment, and install necessary Python dependencies.", justify=False)
    add_code_block("cd backend\npython -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt")
    
    add_normal_paragraph("2. Environment Variable Configuration: Establish a .env file within the backend directory defining cryptographic keys and database connectivity.", justify=False)
    add_code_block("DATABASE_URL=postgresql://postgres:YOUR_PASSWORD@localhost:5432/campus_anpr\nSECRET_KEY=your-secure-key\nALGORITHM=HS256\nACCESS_TOKEN_EXPIRE_MINUTES=1440")
    
    add_normal_paragraph("3. Schema & Data Seeding: Execute the database builder.", justify=False)
    add_code_block("python seed_db.py")
    
    add_normal_paragraph("4. Frontend Initialization: Install React dependencies.", justify=False)
    add_code_block("cd ../frontend\nnpm install")

    # Chapter III
    add_heading('Chapter III. System Execution Procedures', level=1)
    add_normal_paragraph("The application operates via three microservices that must run concurrently.")
    
    add_heading('Automated Execution Sequence:', level=3)
    add_normal_paragraph("Double-click 3_Run_System.bat. The system will dispatch three discrete terminal sessions:")
    add_bullet("Campus Backend Server: Handles logic and API requests (localhost:8000)")
    add_bullet("SMART-PLATE Engine: Conducts OCR inference (localhost:8003)")
    
    add_normal_paragraph("Electron Frontend Startup: The desktop user interface utilizes an orchestration command to bridge the React environment with the Electron desktop shell. When launched, the terminal processes the following script payload from package.json:")
    add_code_block("> campus-anpr-frontend@1.0.0 dev:electron\n> concurrently \"vite\" \"wait-on http://localhost:5173 && node start-electron.cjs\"")
    add_normal_paragraph("This complex execution simultaneously spins up the Vite hot-reloading server, deliberately pauses execution until the port becomes responsive via 'wait-on', and subsequently injects the DOM into the Chromium-based Electron native window (start-electron.cjs) for deployment at the security guardhouse.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: Electron Terminal Executing the Concurrently Command ]")

    # Chapter IV
    document.add_page_break()
    add_heading('Chapter IV. User Roles and Access Credentials', level=1)
    add_normal_paragraph("The system enforces strict Role-Based Access Control (RBAC). Three distinct portals exist to segregate operational duties.")
    
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'User Role'
    hdr_cells[1].text = 'Login Endpoint'
    hdr_cells[2].text = 'Account Email / ID'
    hdr_cells[3].text = 'Account Password'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.paragraph_format.first_line_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    roles_data = [
        ('Administrator', '/admin-login', 'admin@example.com', 'AdminPassword123'),
        ('Security Personnel', '/security-login', 'security@example.com', 'SecurityPassword123'),
        ('Student', '/login', '2024-0001', 'StudentPassword123'),
        ('Faculty', '/login', 'faculty@example.com', 'FacultyPassword123'),
        ('Staff', '/login', 'staff@example.com', 'StaffPassword123'),
        ('Visitor', '/login', 'visitor@example.com', 'VisitorPassword123')
    ]
    
    for role, portal, email, pw in roles_data:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = portal
        row_cells[2].text = email
        row_cells[3].text = pw
        for cell in row_cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Inches(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
    add_normal_paragraph("Standard users operate strictly through the primary unified login portal.")

    # Chapter V
    document.add_page_break()
    add_heading('Chapter V. Core Module Functionality & Navigation', level=1)
    
    add_heading('5.1 Administrator Module', level=2)
    add_normal_paragraph("The Administrative suite provides complete oversight over campus parameters, database records, and hardware configurations.")
    add_bullet("System Dashboard: Presents overarching data visualizations detailing daily entry/exit throughput, active vehicle counts, and violation metrics.")
    add_bullet("User Directory Management: Provision, suspend, or audit accounts for all personnel, inclusive of Security Guards.")
    add_bullet("Vehicle Registration Queues: A dedicated interface to inspect, verify, and authorize vehicle registration payloads submitted by campus constituents.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: Administrator Dashboard Overview ]")
    
    add_heading('5.2 Security Personnel Module', level=2)
    add_normal_paragraph("The Security module is engineered for high-visibility, rapid-response monitoring at the physical campus gates via the Electron Desktop Interface.")
    add_bullet("Real-Time Surveillance Feed: Monitors incoming/outgoing lanes with live video.")
    add_bullet("Intelligent Alert System: Instantly triggers high-contrast Toasts and audible alerts upon detecting blacklisted plates.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: Security Real-Time Monitoring Feed Showing Plate Detection ]")
    
    add_heading('5.3 Standard User Self-Service Module', level=2)
    add_normal_paragraph("A streamlined web interface for constituents (Students, Faculty, Staff) to manage their assets without requiring direct administrative intervention.")
    add_bullet("Asset Registry: Review currently registered vehicles, view digital campus passes, and check their respective administrative approval status.")
    add_bullet("Registration Request Protocol: Submit plate numbers, vehicle classifications, models, and color descriptors to the admin approval queue.")
    
    add_heading('5.4 System Navigation Workflows', level=2)
    add_normal_paragraph("The user interface relies on React Router DOM to facilitate seamless, Single-Page Application (SPA) navigation without triggering full page reloads. The navigation paradigm is highly dependent on the authenticated user's role:")
    add_normal_paragraph("Admin Navigation Hierarchy: Administrators utilize a persistent left-hand collapsing sidebar. Routing transitions instantaneously between /admin/dashboard, /admin/users, and /admin/cameras. The layout preserves state, meaning search filters in the User Directory are maintained even if the admin temporarily clicks into the Settings tab.")
    add_normal_paragraph("Security Dashboard Rigidity: Unlike the Administrator panel, the Security Guard portal is designed as a rigid, monolithic interface (/security/dashboard). To prevent guards from accidentally closing the live camera feed, standard navigation tabs are intentionally suppressed. Overrides and log checks occur within modular modal popups overlaying the live feed.")
    add_screenshot_placeholder("[ INSERT SCREENSHOT: Admin Navigation Sidebar Expanded ]")

    # Chapter VI
    document.add_page_break()
    add_heading('Chapter VI. Hardware & Camera Integration', level=1)
    add_normal_paragraph("The Campus ANPR system acts as the central hub for hardware camera integration, specifically optimized for standard USB webcams.")
    add_bullet("Camera Setup: The system utilizes directly connected USB webcams or integrated laptop webcams. The active webcam feed is processed directly through the security interface without requiring complex IP network configurations.")
    add_bullet("OCR Processing: The SMART-PLATE Engine continuously polls the active webcam video feed and evaluates frames for plate detection at up to 15 frames per second.")

    # Chapter VII
    document.add_page_break()
    add_heading('Chapter VII. Troubleshooting & System Administration', level=1)
    
    add_heading('7.1 Port Collision Constraints', level=2)
    add_normal_paragraph("Symptom: Execution halts due to 'Address already in use' on Ports 8000, 8003, or 5173.", justify=False)
    add_normal_paragraph("Resolution: Terminate orphaned processes by executing the scripts/stop_all.bat utility.", justify=False)
    
    add_heading('7.2 Database Connectivity Failures', level=2)
    add_normal_paragraph("Symptom: The FastAPI backend throws 'psycopg2.OperationalError' or fails to bind to localhost:5432.", justify=False)
    add_normal_paragraph("Resolution: Verify that the PostgreSQL service (postgresql-x64-14) is actively running in the Windows Services Manager (services.msc).", justify=False)

    output_filename = os.path.join('docs', 'USER_GUIDE_Professional.docx')
    try:
        document.save(output_filename)
        print(f"{output_filename} has been created successfully with strict professional layout.")
    except PermissionError:
        output_filename = os.path.join('docs', 'USER_GUIDE_Professional_v3.docx')
        document.save(output_filename)
        print(f"Original file was open. Saved as {output_filename} instead.")

if __name__ == "__main__":
    create_docx()
